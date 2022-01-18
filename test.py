#!/usr/bin/env python3
import docx
from htmldocx import HtmlToDocx
import mammoth

temp = docx.Document("temp.docx")
new_parser = HtmlToDocx()
html = '<p>Le désordre d’une certaine réalisatrice/productrice, avec des effets spéciaux qui ont fait des ravages dans la mémoire collective des acteurs et de l’équipe, avait fait tout le chemin depuis la déclaration d’intention du réalisateur de le terminer uniquement grâce à mes efforts.</p><p>Aujourd’hui, c’était la première des <em><u>Aventures de Mikuru Asahina Épisode&nbsp;00</u></em>, et même si je n’étais pas sûr qu’il s’agisse d’un film ou simplement d’une vidéo de promotion pour Asahina, il était probablement accueilli avec des critiques élogieuses en ce moment dans la salle audiovisuelle.</p>'
docout = new_parser.parse_html_string(html)
out_para = temp.paragraphs[1]
out_para.text = ""
print(len(docout.paragraphs))
for run in docout.paragraphs[1].runs:
    output_run = out_para.add_run(run.text)
    # Run's font data
    output_run.style = run.style
    # Run's color data
    output_run.font.color.rgb = run.font.color.rgb
    # Run's bold data
    output_run.bold = run.bold
    # Run's italic data
    output_run.italic = run.italic
    # Run's underline data
    output_run.underline = run.underline
    out_para.style.name = docout.paragraphs[0].style.name

toshearch = "dffzrzu"
popfile = open("database/pop_french.txt", 'r', encoding="utf-8")
popword = [line.split('\n') for line in popfile.readlines()]
matching = [i for i in popword if i[0].startswith(toshearch)]
if matching == None:
    globalfile = open("database/french.txt", 'r', encoding="utf-8")
    word = [line.split('\n') for line in globalfile.readlines()]
    matching = [i for i in popword if i[0].startswith(toshearch)]
try:
    if matching[0][0] == toshearch:
        try:
            print (matching[1][0])
        except:
            print("null")
    else:
        print(matching[0][0])
except:
    print("null")


temp.save("temp.docx")
docout.save("myfinaltext.docx")



