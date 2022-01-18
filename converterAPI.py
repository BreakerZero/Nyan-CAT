from pprint import pprint

import docx
from htmldocx import HtmlToDocx
import re, argparse


class ConverterAPI:
    def __init__(self):
        self.new_parser = HtmlToDocx()

    def ParaHtmlToDocx(self, Html: str, Docx: docx.document, ParaPosition: int, Offset: int, SaveName: str):
        tempdocx = self.new_parser.parse_html_string(Html)  # docx file with only the html conversion
        for ParaInHtml in range(0, len(tempdocx.paragraphs)):
            if len(tempdocx.paragraphs) + ParaPosition + Offset > len(Docx.paragraphs):
                Docx.add_paragraph()
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].text = ""
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].alignment = tempdocx.paragraphs[ParaInHtml].alignment
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].style = tempdocx.paragraphs[ParaInHtml].style
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].style.name = tempdocx.paragraphs[ParaInHtml].style.name
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].style.base_style = tempdocx.paragraphs[ParaInHtml].style.base_style
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].style.priority = tempdocx.paragraphs[ParaInHtml].style.priority
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].style.style_id = tempdocx.paragraphs[ParaInHtml].style.style_id
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].paragraph_format.left_indent = tempdocx.paragraphs[ParaInHtml].paragraph_format.left_indent
            Docx.paragraphs[ParaPosition + Offset + ParaInHtml].paragraph_format.right_indent = tempdocx.paragraphs[ParaInHtml].paragraph_format.right_indent
            for idx, run in enumerate(tempdocx.paragraphs[ParaInHtml].runs):
                if len(tempdocx.paragraphs[ParaInHtml].runs) > len(Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs):
                    Docx.paragraphs[ParaPosition + Offset + ParaInHtml].add_run()
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].text = run.text
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].bold = run.bold
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].italic = run.italic
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].underline = run.underline
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].font.color.rgb = run.font.color.rgb
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].style.name = run.style.name
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].font.name = run.font.name
                Docx.paragraphs[ParaPosition + Offset + ParaInHtml].runs[idx].font.size = run.font.size
            Docx.save(SaveName)

    def ParaDocxToHtml(self, Docx: docx.document, ParaPosition: int):
        ParatoConvert = Docx.paragraphs[ParaPosition]
        style = ParatoConvert.style.name
        valclass = ""
        if style == "Normal":
            tag = 'p'
        else:
            tag = re.sub(r'[a-z- ]+', '', style).lower()
        if tag == 't':
            tag = 'h1'
        if ParatoConvert.alignment or ParatoConvert.paragraph_format.left_indent:
            valclass = valclass + " class=\""
            if str(ParatoConvert.alignment) == "CENTER (1)":
                valclass = valclass + "ql-align-center"
            elif str(ParatoConvert.alignment) == "JUSTIFY (3)":
                valclass = valclass + "ql-align-justify"
            elif str(ParatoConvert.alignment) == "RIGHT (2)":
                valclass = valclass + "ql-align-right"
        if ParatoConvert.paragraph_format.left_indent and ParatoConvert.alignment:
            valclass = valclass + " "
        if ParatoConvert.paragraph_format.left_indent:
            num = round(float(ParatoConvert.paragraph_format.left_indent.cm)*2.37/3*4/5.5)
            valclass = valclass + "ql-indent-" + str(num)
        if ParatoConvert.paragraph_format.left_indent or ParatoConvert.alignment:
            valclass = valclass + "\""
        Html = '<' + tag + valclass + '>'
        for run in ParatoConvert.runs:
            if run.text != '':
                if run.text[0] == '\n':
                    Html = Html + "</" + tag + "><" + tag + ">"
            if run.font.color.rgb is not None:
                RGB = tuple(int(str(run.font.color.rgb)[i:i + 2], 16) for i in (0, 2, 4))
                Html = Html + "<span style=\"color: rgb(" + str(RGB[0]) + ", " + str(RGB[1]) + ", " + str(RGB[2]) + ");\">"
            if run.italic:
                Html = Html + "<em>"
            if run.bold:
                Html = Html + "<strong>"
            if run.underline:
                Html = Html + "<u>"
            Html = Html + run.text
            if run.underline:
                Html = Html + "/<u>"
            if run.bold:
                Html = Html + "</strong>"
            if run.italic:
                Html = Html + "</em>"
            if run.font.color.rgb is not None:
                Html = Html + "</span>"
        Html = Html + '</' + tag + '>'
        print(Html)
