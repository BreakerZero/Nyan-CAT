# -*- coding: utf-8 -*-
import json
import docx
import csv
import cv2
import numpy as np
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import base64
import io
import os

from backend.converterAPI import ConverterAPI
from flask import Flask, request, jsonify, render_template, redirect, sessions, url_for, flash, abort
import flask_login
from werkzeug.security import generate_password_hash, check_password_hash
import flask_sqlalchemy
from werkzeug.utils import redirect, secure_filename
from backend.translateAPI import TranslatorAPI
from flask_login import UserMixin, LoginManager, login_user, current_user, login_required, logout_user
from flask_dropzone import Dropzone
import html2text
import threading

app = Flask(__name__)
app.config['DROPZONE_REDIRECT_VIEW'] = "home"
app.config['DROPZONE_DEFAULT_MESSAGE'] = "Déposez vos fichiers ici (faites un glisser-deposer ou cliquez pour ouvrir)"
app.config['DROPZONE_ALLOWED_FILE_CUSTOM'] = True
app.config['DROPZONE_ALLOWED_FILE_TYPE'] = '.png, .jpg, .jpeg, .pdf, .docx, .doc, .odt,.txt'
app.config['DROPZONE_INVALID_FILE_TYPE'] = "L'extension de ce fichier de correspond pas avec votre précédente sélection"
app.config['DROPZONE_UPLOAD_MULTIPLE'] = True
translator = TranslatorAPI('./translatemodel/')  # chemin vers les modèles
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
file_path = os.path.abspath(os.getcwd()) + "\\database\\nyan.db"
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + str(file_path)  # Nom de la bdd
app.config['SECRET_KEY'] = '9df31cd3eb2f6f6386571da69d6b418e'  # Clé random pour autentification
app.config['UPLOAD_FOLDER'] = "fileproject"
app.config['UPLOAD_EXTENSIONS'] = ['.png', '.jpg', '.jpeg', '.pdf', '.docx', '.doc', '.odt',
								   '.txt']  # extensions autorisées
app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 50  # max 50Mo
db = flask_sqlalchemy.SQLAlchemy(app)  # lien bdd
app.config["DEBUG"] = True  # option debug
login_manager = LoginManager()
login_manager.login_view = '/login'
login_manager.init_app(app)
dropzone = Dropzone(app)
ConvAPI = ConverterAPI()


@app.errorhandler(413)
def too_large(e):
	return "Fichier trop volumineux", 413


@login_manager.user_loader
def load_user(user_id):
	# since the user_id is just the primary key of our user table, use it in the query for the user
	return db.session.get(User, int(user_id))


def verify_owner(project_id):
	return str(flask_login.current_user.id) == str(Project.query.filter_by(id=project_id).first().Owner)


def manage_csv_memory(project_id):
	csv_path = f"static/csv/memory{project_id}.csv"
	if not os.path.isfile(csv_path):
		with open(csv_path, "w") as csvfile:
			writer = csv.writer(csvfile, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator='\n')
			project = Project.query.filter_by(id=project_id).first()
			writer.writerow([project.Source_Lang, project.Target_Lang])


def get_project_data_for_get_method(project_id):
	project = Project.query.filter_by(id=project_id).first()
	user = User.query.filter_by(id=flask_login.current_user.id).first()
	return (
		project.Type,
		project.Extension,
		project.Last_Block,
		os.listdir(os.path.join(app.config['UPLOAD_FOLDER'], str(project_id))),
		user.KeepStyle,
		user.Autocomplete,
		user.TranslatorSettings,
		project
	)


def get_project_data_for_post_method(project_id):
	project = Project.query.filter_by(id=project_id).first()
	user = User.query.filter_by(id=flask_login.current_user.id).first()
	return (
		project.Type,
		project.Extension,
		project.Source_Lang,
		project.Target_Lang,
		user.TranslatorProvider,
		user.TranslatorSettings,
		user.Formality,
		user.ApiKey
	)


def get_context_paragraphs(i, parasin, direction="before"):
	context = []

	# Set the range for previous (before) or next (after) paragraphs
	if direction == "before":
		start = max(0, i - 5)
		end = i
		step = 1
	else:
		start = i + 1
		end = min(len(parasin), i + 6)
		step = 1

	# Iterate through the range to gather context paragraphs
	for j in range(start, end, step):
		para_text = parasin[j].text.strip()
		if para_text:
			context.append(para_text)

	return " ".join(context)

class Glossary(db.Model):  # Modèle du glossaire
	id = db.Column(db.Integer, primary_key=True, nullable=False, unique=True)
	Source_Lang = db.Column(db.Text, nullable=False)
	Target_Lang = db.Column(db.Text, nullable=False)
	Source = db.Column(db.Text, nullable=False)
	Target = db.Column(db.Text, nullable=False)


class User(UserMixin, db.Model):  # Modèle utilisateur
	id = db.Column(db.Integer, primary_key=True, nullable=False, unique=True)
	Pseudo = db.Column(db.Text, nullable=False, unique=True)
	Mail = db.Column(db.Text, nullable=False, unique=True)
	Password = db.Column(db.Text, nullable=False)
	Status = db.Column(db.Integer)
	TranslatorSettings = db.Column(db.Text, nullable=False)
	TranslatorProvider = db.Column(db.Text, nullable=False)
	Formality = db.Column(db.Text)
	ApiKey = db.Column(db.Text, nullable=False)
	KeepStyle = db.Column(db.Integer)
	Autocomplete = db.Column(db.Integer)


class Project(db.Model):  # Modèle projet
	id = db.Column(db.Integer, primary_key=True, nullable=False, unique=True)
	Name = db.Column(db.Text, nullable=False)
	Type = db.Column(db.Text, nullable=False)
	Owner = db.Column(db.Text, nullable=False)
	Extension = db.Column(db.Text, nullable=False)
	Source_Lang = db.Column(db.Text, nullable=False)
	Target_Lang = db.Column(db.Text, nullable=False)
	Advancement = db.Column(db.Integer, nullable=False)
	Last_Block = db.Column(db.Integer, nullable=False)
	Last_Previous_Block = db.Column(db.Integer, nullable=False)


class Lexicon_fr(db.Model):
	id = db.Column(db.Integer, primary_key=True, nullable=False, unique=True)
	ortho = db.Column(db.String, nullable=False)
	freqfilms = db.Column(db.Float, nullable=False)
	freqlivres = db.Column(db.Float, nullable=False)


class TranslationMemory(db.Model):  # Modèle mémoire de traduction
	id = db.Column(db.Integer, primary_key=True, nullable=False, unique=True)
	Source_Lang = db.Column(db.Text, nullable=False)
	Target_Lang = db.Column(db.Text, nullable=False)
	Source = db.Column(db.Text, nullable=False)
	Target = db.Column(db.Text, nullable=False)
	Owner = db.Column(db.Integer, nullable=False)
	Project = db.Column(db.Integer, nullable=False)
	Segment = db.Column(db.Integer, nullable=False)


def Glos():  # fonction formatage glossaire
	formatedGlo = ""
	i = 0
	Gloquery = Glossary.query.order_by(Glossary.Source, Glossary.Target).with_entities(Glossary.Source,
																					   Glossary.Target).all()
	if Gloquery is None:
		return formatedGlo
	else:
		for line in Gloquery:
			formatedGlo = formatedGlo + str(Gloquery[i][0]) + "\t" + str(Gloquery[i][1]) + "\n"
			i = i + 1
		return formatedGlo[:-1]


with app.app_context():
	formatedGlossary = Glos()


@app.route('/', methods=["GET"])  # racine site
def index():
	return render_template('index.html')


@app.route('/autocomplete', methods=["POST"])
@login_required
def autocomplete():
	tosearch = request.json["begin"]
	tosearch = str(tosearch).lower()
	matching = Lexicon_fr.query.filter(Lexicon_fr.ortho.startswith(tosearch)).order_by(Lexicon_fr.freqlivres.desc(),
																					   Lexicon_fr.freqfilms.desc()).limit(
		2).all()
	if len(matching) == 2:
		if str(matching[0].ortho) == tosearch and str(matching[0].ortho) != str(matching[1].ortho):
			if " " not in str(matching[1].ortho):
				return jsonify({"result": matching[1].ortho})
			else:
				return jsonify({"result": " "})
		elif str(matching[0].ortho) == tosearch and str(matching[0].ortho) == str(matching[1].ortho):
			return jsonify({"result": " "})
		else:
			if " " not in str(matching[0].ortho):
				return jsonify({"result": matching[0].ortho})
			else:
				return jsonify({"result": " "})
	elif len(matching) == 1:
		if str(matching[0].ortho) == tosearch:
			return jsonify({"result": " "})
		else:
			if " " not in str(matching[0].ortho):
				return jsonify({"result": matching[0].ortho})
			else:
				return jsonify({"result": " "})
	else:
		return jsonify({"result": " "})


@app.route('/translate', methods=["POST"])  # chemin de l'API de traduction
@login_required
def get_prediction():
	provider = request.json['provider']
	apikey = request.json['apikey']
	source = request.json['source']
	target = request.json['target']
	formality = request.json['formality']
	text = request.json['text']
	translation = TranslatorAPI.translate(translator, provider, apikey, source, target, formality, text,
										  formatedGlossary)
	return jsonify({"output": translation})


@app.route("/glossary", methods=["GET", "POST"])
@login_required
def glossary():
	if request.method == "POST":
		G_Source_Lang = request.form['Source_Lang']
		G_Target_Lang = request.form['Target_Lang']
		G_Source = request.form['Source']
		G_Target = request.form['Target']
		New_Glossary = Glossary(Source_Lang=G_Source_Lang, Target_Lang=G_Target_Lang, Source=G_Source, Target=G_Target)
		try:
			db.session.add(New_Glossary)
			db.session.commit()
			flash("Ajout dans le glossaire réussie")
			return redirect("/glossary")
		except:
			flash("l'ajout dans le glossaire a échoué")
			return redirect("/glossary")
	else:
		EditGlossary = Glossary.query.order_by(Glossary.id).with_entities(Glossary.id, Glossary.Source_Lang,
																		  Glossary.Target_Lang, Glossary.Source,
																		  Glossary.Target).all()
		return render_template('glossary.html', EditGlossary=EditGlossary)


@app.route("/glossary/delete/<int:id>")
@login_required
def Glossary_Delete(id):
	Glossary_to_Delete = Glossary.query.get_or_404(id)
	try:
		db.session.delete(Glossary_to_Delete)
		db.session.commit()
		flash("Suppression Réussie")
		return redirect("/glossary")
	except:
		flash("La suppression dans le glossaire a échoué")
		return redirect("/glossary")


@app.route("/glossary/update/<int:id>", methods=["GET", "POST"])
@login_required
def Glossary_Update(id):
	Glossary_to_Update = Glossary.query.get_or_404(id)
	if request.method == "POST":
		Glossary_to_Update.Source_Lang = request.form['Source_Lang']
		Glossary_to_Update.Target_Lang = request.form['Target_Lang']
		Glossary_to_Update.Source = request.form['Source']
		Glossary_to_Update.Target = request.form['Target']
		try:
			db.session.commit()
			return redirect("/glossary")
		except:
			return "la mise à jour dans le glossaire a échoué"
	else:
		return render_template('update.html', Glossary_to_Update=Glossary_to_Update)


@app.route('/login', methods=["GET", "POST"])
def login():
	if request.method == "POST":
		pseudo = request.form.get('pseudo')
		password = request.form.get('password')
		remember = True if request.form.get('remember') else False
		testpseudo = User.query.filter_by(Pseudo=pseudo).first()
		if not testpseudo or not check_password_hash(testpseudo.Password, password):
			flash("Authentification impossible, vérifiez vos informations d'identification.")
			return redirect("/login")
		else:
			jsonpath = "static/json/memory" + str(testpseudo.id) + ".json"
			jsonfile = open(jsonpath, "w", encoding="UTF-8")
			data = TranslationMemory.query.filter_by(Owner=int(testpseudo.id)).all()
			jsondata = []
			for i in data:
				jsondata.append(
					{"id": i.id, "Source_Lang": i.Source_Lang, "Target_Lang": i.Target_Lang, "Source": i.Source,
					 "Target": i.Target})
			json.dump(jsondata, jsonfile)
			jsonfile.close()
			login_user(testpseudo, remember=remember)
			return redirect("/home")
	else:
		return render_template('login.html')


@app.route('/signup', methods=["GET", "POST"])
def signup():
	if request.method == "POST":
		mail = request.form.get('mail')
		pseudo = request.form.get('pseudo')
		password = request.form.get('password')
		testmail = User.query.filter_by(Mail=mail).first()  # verification pseudo déjà existant
		if testmail:  # si utilisateur déjà existant retour à la page de connexion en disant que l'adresse existe déjà
			flash("Cette adresse est déjà utilisée:")
			return redirect("/signup")
		testpseudo = User.query.filter_by(Pseudo=pseudo).first()
		if testpseudo:
			flash("Ce nom d'utilisateur est déjà utilisé, choisissez-en un autre ou :")
			return redirect("/signup")
		hashpassword = generate_password_hash(password, method="scrypt")

		new_user = User(Pseudo=pseudo, Mail=mail, Password=hashpassword, Status=0, TranslatorSettings="More",
						TranslatorProvider="Nyan-Cat", ApiKey="None", KeepStyle=0, Autocomplete=0)
		# add the new user to the database
		db.session.add(new_user)
		db.session.commit()
		return redirect("/login")
	else:
		return render_template('signup.html')


@app.route('/logout')
@login_required
def logout():
	if os.path.exists("static/json/memory" + str(flask_login.current_user.id) + ".json"):
		os.remove("static/json/memory" + str(flask_login.current_user.id) + ".json")
	logout_user()
	return redirect("/")


@app.route('/home')
@login_required
def home():
	current = str(flask_login.current_user.id)
	Projectlist = Project.query.filter_by(Owner=current).all()
	return render_template('home.html', pseudo=current_user.Pseudo, Projectlist=Projectlist)


@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings():
	currentid = str(flask_login.current_user.id)
	User_to_Update = User.query.get_or_404(currentid)
	if request.method == "POST":
		User_to_Update.TranslatorProvider = request.form.get('translatorprovider')
		User_to_Update.ApiKey = request.form.get('api')
		User_to_Update.KeepStyle = True if request.form.get('keepstyle') else False
		User_to_Update.Autocomplete = True if request.form.get('autocomplete') else False
		translatorsettings = request.form.get('translatorsettings')
		if translatorsettings == "Actif":
			User_to_Update.TranslatorSettings = "More"
		elif translatorsettings == "Passif":
			User_to_Update.TranslatorSettings = "Less"
		else:
			User_to_Update.TranslatorSettings = "Disabled"
		formality = request.form.get('formality')
		if formality == "Informel":
			User_to_Update.Formality = "informal"
		elif formality == "Formel":
			User_to_Update.Formality = "formal"
		else:
			User_to_Update.Formality = None
		try:
			db.session.commit()
			return redirect("/home")
		except:
			return "la mise à jour de l'utilisateur a échoué"
	else:
		return render_template('settings.html', User_to_Update=User_to_Update)


@app.route("/newproject", methods=["GET", "POST"])
@login_required
def newproject():
	if request.method == "POST":
		my_files = request.files
		name = request.form.get('name')
		type = request.form.get('type')
		format = request.form.get('format')
		source = request.form.get('source')
		target = request.form.get('target')
		current_owner = str(flask_login.current_user.id)
		try:
			idproject = db.session.query(Project).order_by(Project.id.desc()).first().id
		except:
			idproject = "0"
		idproject = str(int(idproject) + 1)
		os.mkdir(app.config['UPLOAD_FOLDER'] + "/" + idproject)
		for item in my_files:
			up_file = my_files.get(item)
			up_file.filename = secure_filename(up_file.filename)
			if up_file.filename != '':
				file_ext = os.path.splitext(up_file.filename)[1]
			if file_ext not in app.config['UPLOAD_EXTENSIONS']:
				abort(400)
			up_file.save(os.path.join(app.config['UPLOAD_FOLDER'] + "/" + idproject, up_file.filename))
		if type == "Roman/Light Novel (Textuel)":
			type = "text"
		elif type == "Manga/BD (Image)":
			type = "image"
		new_project = Project(id=int(idproject), Name=name, Type=type, Owner=current_owner, Extension=format,
							  Source_Lang=source, Target_Lang=target, Advancement=0, Last_Block=0)
		db.session.add(new_project)
		db.session.commit()
		return redirect('/home')
	else:
		return render_template('newproject.html')


def docx_get_ressource_in_request(request, id, namefile):
	idblock = int(request.json["ressource"])
	OriginalDocx = docx.Document(app.config['UPLOAD_FOLDER'] + "/" + str(id) + "/" + namefile)
	Html = ConverterAPI.ParaDocxToHtml(ConvAPI, OriginalDocx, idblock)
	PreviousHtml = ConverterAPI.ParaDocxToHtml(ConvAPI, OriginalDocx, idblock - 1)
	if PreviousHtml == "<p><br></p>":
		PreviousHtml = "<p></p>"
	NextHtml = ConverterAPI.ParaDocxToHtml(ConvAPI, OriginalDocx, idblock + 1)
	if NextHtml == "<p><br></p>":
		NextHtml = "<p></p>"
	return Html, PreviousHtml, NextHtml


@app.route("/project/text/docx/<int:id>", methods=["GET", "POST"])
@login_required
def projecttextdocx(id):
	if not verify_owner(id):
		logout_user()
		return redirect('/')
	if request.method == "GET":
		#create csv file if not exist
		manage_csv_memory(id)
		#get project data
		(type, extension, last, files, keepstyle, complete, translatorsettings, project) = get_project_data_for_get_method(id)
		if extension == "docx":
			return render_template('docxproject.html', id=id, last=last, keepstyle=keepstyle, complete=complete,
								   user=flask_login.current_user, project=project,
								   translatorsettings=translatorsettings, extension=extension, type=type)
		else:
			return render_template('404.html'), 404
	if request.method == "POST":
		#get project data
		(type, extension, source, target, provider, settings, formality, apikey) = get_project_data_for_post_method(id)
		if extension == "docx":
			files = os.listdir(os.path.join(app.config['UPLOAD_FOLDER'] + "/" + str(id)))
			namefile = str(files[0])
			if "ressource" in request.json:  #Demande ressouce fichier original
				Html, PreviousHtml, NextHtml = docx_get_ressource_in_request(request, id, namefile)
				return jsonify({"result": Html, "previous": PreviousHtml, "next": NextHtml})
			elif "translated" in request.json:  #Demande ressouce fichier traduit
				idblock = int(request.json["translated"])
				idpreviousblock = request.json["previoustranslated"]
				OriginalHtml = str(request.json["originaltext"])
				TranslatedHtml = str(request.json["translatedtext"])
				if TranslatedHtml == "<p><br></p>" or TranslatedHtml == "<p class=""><br></p>":
					TranslatedHtml = "<p></p>"
				SaveName = app.config['UPLOAD_FOLDER'] + "/" + str(id) + "/" + "translated-" + namefile
				if len(files) == 1:  #Création du fichier de sortie s'il n'existe pas
					TranslatedDocx = docx.Document()
					TranslatedDocx.save(SaveName)
				mutex = threading.Lock()
				mutex.acquire()
				while mutex.locked():  #pour gérer un grand nombre d'écriture fichier
					try:
						TranslatedDocx = docx.Document(SaveName)
					except:
						pass
					else:
						mutex.release()
				text = html2text.html2text(TranslatedHtml)
				if idblock < len(TranslatedDocx.paragraphs):  #L'écriture du fichier de sortie s'écrivant au fur et à mesure
					if TranslatedDocx.paragraphs[idblock].text != text:  #Si le text est différent on enregistre ce nouveau texte
						if idpreviousblock is not None:
							ConverterAPI.ParaHtmlToDocx(ConvAPI, TranslatedHtml, TranslatedDocx, int(idpreviousblock), SaveName)
						Html = ConverterAPI.ParaDocxToHtml(ConvAPI, TranslatedDocx, idblock)
					else:
						Html = ConverterAPI.ParaDocxToHtml(ConvAPI, TranslatedDocx, idblock)
				else:  #Si le block n'existe pas dans le document de sortie
					text = html2text.html2text(OriginalHtml)
					parasin = docx.Document(app.config['UPLOAD_FOLDER'] + "/" + str(id) + "/" + namefile).paragraphs
					prev_paragraph = get_context_paragraphs(idblock, parasin, direction="before")
					next_paragraph = get_context_paragraphs(idblock, parasin, direction="after")
					translation = TranslatorAPI.translate(translator, provider, settings, apikey, source, target, formality, text, formatedGlossary, prev_paragraph, next_paragraph)  #on tente une traduction si on a rien dans la mémoire de traduction
					translation = translation.replace("\n", "")
					Html = '<p>' + translation + "</p>"
					ConverterAPI.ParaHtmlToDocx(ConvAPI, Html, TranslatedDocx, idblock, SaveName)
				Project.query.filter_by(id=id).first().Last_Previous_Block = Project.query.filter_by(
					id=id).first().Last_Block
				Project.query.filter_by(id=id).first().Last_Block = idblock
				db.session.commit()
				PreviousHtml = ConverterAPI.ParaDocxToHtml(ConvAPI, TranslatedDocx, idblock - 1)
				if PreviousHtml == "<p><br></p>":
					PreviousHtml = "<p></p>"
				NextHtml = ConverterAPI.ParaDocxToHtml(ConvAPI, TranslatedDocx, idblock + 1)
				if NextHtml == "<p><br></p>":
					NextHtml = "<p></p>"
				return jsonify({"result": Html, "previous": PreviousHtml, "next": NextHtml})


@app.route("/project/text/txt/<int:id>", methods=["GET", "POST"])
@login_required
def projecttexttxt(id):
	return "not implemented yet"


@app.route("/project/text/pdf/<int:id>", methods=["GET", "POST"])
@login_required
def projecttextpdf(id):
	return "not implemented yet"


@app.route("/project/image/png/<int:id>", methods=["GET", "POST"])
@login_required
def projectimagepng(id):
	return "not implemented yet"


@app.route("/project/image/jpg/<int:id>", methods=["GET", "POST"])
@login_required
def projectimagejpg(id):
	return "not implemented yet"


@app.route("/project/image/pdf/<int:id>", methods=["GET", "POST"])
@login_required
def projectimagepdf(id):
	return "not implemented yet"


@app.route("/addsegment", methods=['POST'])
@login_required
def addsegment():
	if request.method == 'POST':
		User_ID = request.json['User_ID']
		Project_ID = request.json['Project_ID']
		Segment_ID = request.json['Segment_ID']
		Source = request.json['Source']
		Target = request.json['Target']
		Source_Lang = request.json['Source_Lang']
		Target_Lang = request.json['Target_Lang']
		#search if the segment already exists
		if TranslationMemory.query.filter_by(Owner=User_ID, Project=Project_ID, Segment=Segment_ID).first() is None:
			New_TranslationMemory = TranslationMemory(Owner=User_ID, Project=Project_ID, Segment=Segment_ID,
													  Source=Source, Target=Target, Source_Lang=Source_Lang,
													  Target_Lang=Target_Lang)
			db.session.add(New_TranslationMemory)
			db.session.commit()
			if os.path.isfile("static/csv/memory" + str(Project_ID) + ".csv"):
				with open("static/csv/memory" + str(Project_ID) + ".csv", "a+") as csvfile:
					writer = csv.writer(csvfile, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL,
										lineterminator='\n')
					Source = Source.replace('\n', '')
					Target = Target.replace('\n', '')
					writer.writerow([Source, Target])
		else:
			TranslationMemory.query.filter_by(Owner=User_ID, Project=Project_ID, Segment=Segment_ID).update(
				{'Source': Source, 'Target': Target, 'Source_Lang': Source_Lang, 'Target_Lang': Target_Lang})
			db.session.commit()
			if os.path.isfile("static/csv/memory" + str(Project_ID) + ".csv"):
				with open("static/csv/memory" + str(Project_ID) + ".csv", "a+") as csvfile:
					writer = csv.writer(csvfile, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL,
										lineterminator='\n')
					Source = Source.replace('\n', '')
					Target = Target.replace('\n', '')
					writer.writerow([Source, Target])
		return jsonify({"result": "ok"})


@app.route("/project/<int:id>/train", methods=["POST"])
@login_required
def train(id):
	"""if os.path.isfile("static/csv/memory" + str(id) + ".csv"):
		with open("static/csv/memory" + str(id) + ".csv", "r") as csvfile:
			reader = csv.reader(csvfile, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator='\n')
			lines = len(list(reader))
		if lines > 1:
			Source_Lang = Project.query.filter_by(id=id).first().Source_Lang
			Target_Lang = Project.query.filter_by(id=id).first().Target_Lang
			#launch the training in new thread
			thread = threading.Thread(target=Training, args=(id, Source_Lang, Target_Lang))
			thread.start()"""
	return jsonify({"result": "ok"})


@app.route('/clone', methods=['POST'])
def clone():
	data = request.json
	img_data = data['image']
	clone_pt = data['clone_pt']
	target_pt = data['target_pt']
	radius = int(data['radius'])

	# Convertir l'image de base64 en image OpenCV
	img_bytes = io.BytesIO(base64.b64decode(img_data))
	pil_image = Image.open(img_bytes)
	image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
	clone = image.copy()

	# Appliquer le clonage avec mélange des bords
	apply_clone_with_blending(image, clone, clone_pt, target_pt, radius)

	# Convertir l'image modifiée en base64 pour la réponse
	_, buffer = cv2.imencode('.jpg', clone)
	img_str = base64.b64encode(buffer).decode('utf-8')

	return jsonify({'image': img_str})


def apply_clone_with_blending(image, clone, clone_pt, target_pt, radius):
	mask = np.zeros((radius * 2, radius * 2), dtype=np.uint8)
	cv2.circle(mask, (radius, radius), radius, 255, -1)
	# Augmenter le flou gaussien pour adoucir davantage les bords
	mask = cv2.GaussianBlur(mask, (31, 31), 0)

	for y in range(-radius, radius):
		for x in range(-radius, radius):
			if x ** 2 + y ** 2 <= radius ** 2:
				src_x = int(clone_pt['x']) + x
				src_y = int(clone_pt['y']) + y
				tgt_x = int(target_pt['x']) + x
				tgt_y = int(target_pt['y']) + y
				if (0 <= src_x < image.shape[1] and 0 <= src_y < image.shape[0] and
						0 <= tgt_x < clone.shape[1] and 0 <= tgt_y < clone.shape[0]):
					alpha = mask[y + radius, x + radius] / 255.0
					clone[tgt_y, tgt_x] = (1.0 - alpha) * image[tgt_y, tgt_x] + alpha * image[src_y, src_x]


ConvAPI = ConverterAPI()  # Utilisation de ton ConverterAPI existant

@app.route('/saveimg/<int:id>', methods=['POST'])
@login_required
def saveimg(id):
	if request.method == 'POST':
		data = request.json
		section = data.get('section')
		imgisbefore = data.get('imgIsBefore')

		# Récupérer l'image base64 et la convertir en une image Pillow
		image_data = data.get('image')
		image_bytes = base64.b64decode(image_data.split(',')[1])
		image = Image.open(io.BytesIO(image_bytes))

		# Créer un objet ImageDraw pour dessiner sur l'image
		draw = ImageDraw.Draw(image)

		# Récupérer les blocs de texte et les ajouter sur l'image
		text_blocks = data.get('textBlocks', [])
		line_height_factor = 1.5  # Facteur de hauteur de ligne pour correspondre à la `line-height` de l'interface web
		for block in text_blocks:
			content = block['content']
			x = block['x']
			y = block['y']
			width = block['width']
			height = block['height']
			sizefont = block['sizefont']

			# Utiliser une police par défaut (ou charger une autre police si nécessaire)
			try:
				font_path = "arial.ttf"
				font = ImageFont.truetype(font_path, size=int(sizefont))
			except IOError:
				font = ImageFont.load_default()

			# Diviser le texte en lignes en respectant les retours à la ligne explicites
			wrapped_lines = []
			for line in content.split('\n'):
				current_line = ""
				for word in line.split():
					# Ajouter chaque mot à la ligne actuelle
					test_line = current_line + (" " if current_line else "") + word
					# Calculer la largeur de la ligne actuelle si le mot est ajouté
					line_width = draw.textbbox((0, 0), test_line, font=font)[2]
					# Si la largeur dépasse le cadre, sauvegarder la ligne actuelle et commencer une nouvelle
					if line_width > width:
						wrapped_lines.append(current_line)
						current_line = word
					else:
						current_line = test_line

				# Ajouter la dernière ligne de chaque paragraphe (segment séparé par '\n')
				if current_line:
					wrapped_lines.append(current_line)
				# Ajouter une ligne vide après un retour à la ligne explicite
				wrapped_lines.append("")

			# Calculer la position pour commencer à dessiner le texte
			text_x = x
			text_y = y

			# Dessiner chaque ligne de texte dans le cadre avec la `line-height` de 1.1
			for line in wrapped_lines:
				if line.strip() != "":
					line_bbox = draw.textbbox((0, 0), line, font=font)
					line_width = line_bbox[2] - line_bbox[0]
					# Dessiner le texte centré horizontalement à l'intérieur du cadre
					draw.text((text_x + (width - line_width) / 2, text_y), line, fill="black", font=font)
					# Déplacer le curseur de dessin vers la ligne suivante en utilisant `line-height`
					text_y += (line_bbox[3] - line_bbox[1]) * line_height_factor

		# Convertir l'image Pillow en base64 pour l'utiliser dans le document Word
		buffered = io.BytesIO()
		image.save(buffered, format="PNG")
		img_str = base64.b64encode(buffered.getvalue()).decode()

		# Charger ou créer un document Word
		files = os.listdir(os.path.join(app.config['UPLOAD_FOLDER'] + "/" + str(id)))
		targetfile = str(files[1])

		if targetfile:
			targetdoc = Document(os.path.join(app.config['UPLOAD_FOLDER'], str(id), targetfile))
		else:
			return render_template('404.html'), 404

		# Insérer l'image modifiée dans le document Word avec ton ConverterAPI
		ConvAPI.insert_image_from_base64(targetdoc, img_str, ParaPosition=int(section), imgisbefore=imgisbefore)

		# Sauvegarder le document Word
		targetdoc.save(os.path.join(app.config['UPLOAD_FOLDER'], str(id), targetfile))

		return redirect("/project/text/docx/" + str(id))


app.run(host="127.0.0.1", port=5000, threaded=True)
